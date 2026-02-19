#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Costanti (EMU)
EMU_PER_INCH = 914400

# Dimensioni richieste (già in EMU come da tuo XML)
CX = 6336792   # 6.93" * 914400
CY = 1581912   # 1.73" * 914400

# Offset richiesti
POS_X = 0
POS_Y = 129857  # come nel frammento "corretto"

# relativeHeight richiesto
REL_HEIGHT = 251661312


def _next_docpr_id(doc: Document) -> int:
    """
    Trova un id libero per wp:docPr evitando collisioni.
    """
    ids = set()
    # scandisce tutti i docPr presenti nel documento
    for el in doc.element.xpath(".//wp:docPr"):
        v = el.get("id")
        if v is not None:
            try:
                ids.add(int(v))
            except ValueError:
                pass
    i = 1
    while i in ids:
        i += 1
    return i


def add_body_background_anchor(paragraph, image_path: str, docpr_id: int):
    """
    Inserisce l'immagine come oggetto floating (wp:anchor) nel BODY,
    behind text, posizionata:
      - H: relativeFrom="column", offset 0
      - V: relativeFrom="paragraph", offset 129857
    con dimensioni CX/CY.
    """
    # Crea la relationship per l'immagine e ottiene rId
    rId, _ = paragraph.part.get_or_add_image(image_path)

    # nsdecls doesn't include the Word 2010 "wp14" prefix by default,
    # so append its declaration explicitly.
    ns = nsdecls('w', 'r', 'wp', 'a', 'pic') + ' xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"'

    anchor_xml = f"""
    <w:r {ns}>
      <w:drawing>
        <wp:anchor distT="0" distB="0" distL="0" distR="0"
                   simplePos="0"
                   relativeHeight="{REL_HEIGHT}"
                   behindDoc="1"
                   locked="0"
                   layoutInCell="1"
                   allowOverlap="1">

          <wp:simplePos x="0" y="0"/>

          <wp:positionH relativeFrom="column">
            <wp:posOffset>{POS_X}</wp:posOffset>
          </wp:positionH>

          <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>{POS_Y}</wp:posOffset>
          </wp:positionV>

          <wp:extent cx="{CX}" cy="{CY}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>

          <wp:wrapNone/>

          <wp:docPr id="{docpr_id}" name="BodyBackground"/>
          <wp:cNvGraphicFramePr/>

          <a:graphic>
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:pic>
                <pic:nvPicPr>
                  <pic:cNvPr id="0" name="BackgroundImage"/>
                  <pic:cNvPicPr/>
                </pic:nvPicPr>

                <pic:blipFill>
                  <a:blip r:embed="{rId}"/>
                  <a:stretch>
                    <a:fillRect/>
                  </a:stretch>
                </pic:blipFill>

                <pic:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="{CX}" cy="{CY}"/>
                  </a:xfrm>
                  <a:prstGeom prst="rect">
                    <a:avLst/>
                  </a:prstGeom>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphic>

          <wp14:sizeRelH relativeFrom="margin">
            <wp14:pctWidth>0</wp14:pctWidth>
          </wp14:sizeRelH>
          <wp14:sizeRelV relativeFrom="margin">
            <wp14:pctHeight>0</wp14:pctHeight>
          </wp14:sizeRelV>

        </wp:anchor>
      </w:drawing>
    </w:r>
    """

    paragraph._p.append(parse_xml(anchor_xml))


def main():
    ap = argparse.ArgumentParser(
        description="Inserisce un'immagine floating (behind text) nel BODY con posizionamento column/paragraph."
    )
    ap.add_argument("--template", required=True, help="DOCX template di input.")
    ap.add_argument("--image", required=True, help="Immagine da inserire come sfondo (png/jpg).")
    ap.add_argument("--out", required=True, help="DOCX di output.")
    ap.add_argument("--paragraph-index", type=int, default=0,
                    help="Indice del paragrafo in cui inserire l'anchor (default: 0).")
    args = ap.parse_args()

    doc = Document(args.template)

    # Paragrafo target (default: primo)
    if doc.paragraphs:
        idx = max(0, min(args.paragraph_index, len(doc.paragraphs) - 1))
        p = doc.paragraphs[idx]
    else:
        p = doc.add_paragraph("")

    docpr_id = _next_docpr_id(doc)
    add_body_background_anchor(p, args.image, docpr_id)

    doc.save(args.out)


if __name__ == "__main__":
    main()
